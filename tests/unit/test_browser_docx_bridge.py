from __future__ import annotations

import base64
import json
import threading
from http.client import HTTPConnection
from pathlib import Path

from docx import Document

from ai_pm_lab_privacy_gate.application.privacy_service import PrivacyGateService
from ai_pm_lab_privacy_gate.domain.models import Finding, PageContent
from ai_pm_lab_privacy_gate.infrastructure.local_api.browser_ai_persistence import (
    install_browser_ai_persistence,
)
from ai_pm_lab_privacy_gate.infrastructure.local_api.browser_docx import (
    install_browser_docx_support,
)
from ai_pm_lab_privacy_gate.infrastructure.local_api.browser_pairing import BrowserPairingRegistry
from ai_pm_lab_privacy_gate.infrastructure.local_api.server import create_local_api_server
from ai_pm_lab_privacy_gate.infrastructure.security.secret_store import MemorySecretStore
from ai_pm_lab_privacy_gate.infrastructure.storage.ai_library_repository import AiLibraryRepository


EMAIL = "jane.smith@example.com"
ORIGIN = "chrome-extension://privacygate-docx-test"
AUTH_TOKEN = "test-local-docx-auth-token-0123456789"


class FakePiiEngine:
    document_language = "en"

    def analyze_page(self, page: PageContent, _profile) -> list[Finding]:
        start = page.text.find(EMAIL)
        if start < 0:
            return []
        end = start + len(EMAIL)
        return [
            Finding(
                finding_id=f"p{page.page_number}-{start}-{end}-0",
                entity_type="EMAIL_ADDRESS",
                text=EMAIL,
                start=start,
                end=end,
                score=0.99,
                page_number=page.page_number,
                context=page.text,
            )
        ]


def _request(server, method: str, path: str, payload=None, *, token=None, origin=ORIGIN):
    connection = HTTPConnection("127.0.0.1", server.server_port, timeout=3)
    body = None if payload is None else json.dumps(payload)
    headers = {}
    if payload is not None:
        headers["Content-Type"] = "application/json"
    if token:
        headers["Authorization"] = f"Bearer {token}"
    if origin:
        headers["Origin"] = origin
    connection.request(method, path, body=body, headers=headers)
    response = connection.getresponse()
    raw = response.read().decode("utf-8")
    connection.close()
    return response.status, json.loads(raw) if raw else {}


def _source_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("Contact ")
    first.bold = True
    paragraph.add_run(EMAIL)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Project reference only"
    document.core_properties.author = "Original Author"
    document.save(path)
    return path


def _pair(server) -> str:
    challenge = server.browser_pairing.create_challenge()
    status, payload = _request(
        server,
        "POST",
        "/v1/browser/pair",
        {"code": challenge.code, "client_name": "docx-unit-test"},
    )
    assert status == 200
    return payload["browser_token"]


def test_docx_browser_round_trip_preserves_format_and_removes_original(tmp_path: Path) -> None:
    service = PrivacyGateService(pii_engine=FakePiiEngine())
    pairing = BrowserPairingRegistry(MemorySecretStore())
    server = create_local_api_server(
        service,
        port=0,
        auth_token=AUTH_TOKEN,
        browser_pairing=pairing,
    )
    repository = AiLibraryRepository(tmp_path / "library")
    assert install_browser_ai_persistence(server, repository)
    assert install_browser_docx_support(server)
    thread = threading.Thread(target=server.serve_forever, kwargs={"poll_interval": 0.01}, daemon=True)
    thread.start()
    try:
        source = _source_docx(tmp_path)
        token = _pair(server)
        encoded = base64.b64encode(source.read_bytes()).decode("ascii")

        status, analyzed = _request(
            server,
            "POST",
            "/v1/browser/docx/analyze",
            {
                "filename": source.name,
                "file_base64": encoded,
                "profile_key": "general_business",
                "language": "en",
            },
            token=token,
        )
        assert status == 200
        assert analyzed["findings_count"] == 1
        finding = analyzed["findings"][0]
        assert finding["display_value"] == EMAIL
        assert analyzed["review_values_local_only"] is True

        status, protected = _request(
            server,
            "POST",
            "/v1/browser/docx/protect",
            {
                "analysis_id": analyzed["analysis_id"],
                "finding_ids": [finding["finding_id"]],
                "provider": "chatgpt",
            },
            token=token,
        )
        assert status == 200
        assert protected["protected_filename"].endswith("_PrivacyGate.docx")
        assert protected["session_id"]

        snapshot = repository.load_session(protected["session_id"])
        assert snapshot is not None
        assert snapshot.provider == "chatgpt"

        output = tmp_path / "protected.docx"
        output.write_bytes(base64.b64decode(protected["protected_file_base64"]))
        opened = Document(output)
        all_text = "\n".join(
            [paragraph.text for paragraph in opened.paragraphs]
            + [cell.text for table in opened.tables for row in table.rows for cell in row.cells]
        )
        assert EMAIL not in all_text
        assert "[[PG_" in all_text
        assert opened.paragraphs[0].runs[0].bold is True
        assert opened.core_properties.author == "AI PM LAB Privacy Gate"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)


def test_docx_routes_require_browser_pairing(tmp_path: Path) -> None:
    service = PrivacyGateService(pii_engine=FakePiiEngine())
    server = create_local_api_server(service, port=0, auth_token=AUTH_TOKEN)
    assert install_browser_docx_support(server)
    thread = threading.Thread(target=server.serve_forever, kwargs={"poll_interval": 0.01}, daemon=True)
    thread.start()
    try:
        source = _source_docx(tmp_path)
        status, payload = _request(
            server,
            "POST",
            "/v1/browser/docx/analyze",
            {
                "filename": source.name,
                "file_base64": base64.b64encode(source.read_bytes()).decode("ascii"),
                "profile_key": "general_business",
                "language": "en",
            },
        )
        assert status == 401
        assert payload["error"] == "browser_pairing_required"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)
