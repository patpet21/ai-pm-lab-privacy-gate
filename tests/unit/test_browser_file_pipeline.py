from __future__ import annotations

import base64
import json
import threading
from http.client import HTTPConnection
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from ai_pm_lab_privacy_gate.application.privacy_service import PrivacyGateService
from ai_pm_lab_privacy_gate.domain.models import Finding, PageContent
from ai_pm_lab_privacy_gate.infrastructure.local_api.browser_ai_persistence import (
    install_browser_ai_persistence,
)
from ai_pm_lab_privacy_gate.infrastructure.local_api.browser_file import (
    install_browser_file_support,
)
from ai_pm_lab_privacy_gate.infrastructure.local_api.browser_pairing import BrowserPairingRegistry
from ai_pm_lab_privacy_gate.infrastructure.local_api.server import create_local_api_server
from ai_pm_lab_privacy_gate.infrastructure.security.secret_store import MemorySecretStore
from ai_pm_lab_privacy_gate.infrastructure.storage.ai_library_repository import AiLibraryRepository


EMAIL = "jane.smith@example.com"
ORIGIN = "chrome-extension://privacygate-file-test"
AUTH_TOKEN = "test-local-file-auth-token-0123456789"


class FakePiiEngine:
    document_language = "en"

    def analyze_page(self, page: PageContent, _profile) -> list[Finding]:
        start = page.text.find(EMAIL)
        if start < 0:
            return []
        end = start + len(EMAIL)
        return [
            Finding(
                finding_id=f"p{page.page_number}-{start}-{end}-0",
                entity_type="EMAIL_ADDRESS",
                text=EMAIL,
                start=start,
                end=end,
                score=0.99,
                page_number=page.page_number,
                context=page.text,
            )
        ]


def _request(server, path: str, payload: dict, *, token: str | None = None):
    connection = HTTPConnection("127.0.0.1", server.server_port, timeout=20)
    headers = {"Content-Type": "application/json", "Origin": ORIGIN}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    connection.request("POST", path, body=json.dumps(payload), headers=headers)
    response = connection.getresponse()
    raw = response.read().decode("utf-8")
    connection.close()
    return response.status, json.loads(raw) if raw else {}


def _pair(server) -> str:
    challenge = server.browser_pairing.create_challenge()
    status, payload = _request(
        server,
        "/v1/browser/pair",
        {"code": challenge.code, "client_name": "unified-file-unit-test"},
    )
    assert status == 200
    return payload["browser_token"]


def _server(tmp_path: Path):
    service = PrivacyGateService(pii_engine=FakePiiEngine())
    pairing = BrowserPairingRegistry(MemorySecretStore())
    server = create_local_api_server(
        service,
        port=0,
        auth_token=AUTH_TOKEN,
        browser_pairing=pairing,
    )
    assert install_browser_ai_persistence(server, AiLibraryRepository(tmp_path / "library"))
    assert install_browser_file_support(server)
    thread = threading.Thread(target=server.serve_forever, kwargs={"poll_interval": 0.01}, daemon=True)
    thread.start()
    return server, thread


def _protect_round_trip(server, source: Path, token: str) -> tuple[dict, bytes]:
    encoded = base64.b64encode(source.read_bytes()).decode("ascii")
    status, analyzed = _request(
        server,
        "/v1/browser/file/analyze",
        {
            "filename": source.name,
            "file_base64": encoded,
            "profile_key": "general_business",
            "language": "en",
        },
        token=token,
    )
    assert status == 200
    assert analyzed["findings_count"] == 1
    assert analyzed["findings"][0]["display_value"] == EMAIL

    status, protected = _request(
        server,
        "/v1/browser/file/protect",
        {
            "analysis_id": analyzed["analysis_id"],
            "finding_ids": [analyzed["findings"][0]["finding_id"]],
            "provider": "chatgpt",
        },
        token=token,
    )
    assert status == 200
    assert protected["session_id"]
    return protected, base64.b64decode(protected["protected_file_base64"])


def test_unified_file_route_protects_docx(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph(f"Contact {EMAIL}")
    document.save(source)

    server, thread = _server(tmp_path)
    try:
        token = _pair(server)
        protected, raw = _protect_round_trip(server, source, token)
        assert protected["protected_filename"].endswith("_PrivacyGate.docx")
        output = tmp_path / "protected.docx"
        output.write_bytes(raw)
        text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
        assert EMAIL not in text
        assert "[[PG_" in text
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)


def test_unified_file_route_protects_xlsx(tmp_path: Path) -> None:
    source = tmp_path / "sample.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet["A1"] = "Contact"
    worksheet["B1"] = EMAIL
    workbook.save(source)
    workbook.close()

    server, thread = _server(tmp_path)
    try:
        token = _pair(server)
        protected, raw = _protect_round_trip(server, source, token)
        assert protected["protected_filename"].endswith("_PrivacyGate.xlsx")
        output = tmp_path / "protected.xlsx"
        output.write_bytes(raw)
        opened = load_workbook(output, data_only=False)
        try:
            text = "\n".join(
                str(cell.value)
                for sheet in opened.worksheets
                for row in sheet.iter_rows()
                for cell in row
                if cell.value is not None
            )
        finally:
            opened.close()
        assert EMAIL not in text
        assert "[[PG_" in text
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)


def test_unified_file_route_requires_pairing(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph(f"Contact {EMAIL}")
    document.save(source)

    server, thread = _server(tmp_path)
    try:
        status, payload = _request(
            server,
            "/v1/browser/file/analyze",
            {
                "filename": source.name,
                "file_base64": base64.b64encode(source.read_bytes()).decode("ascii"),
                "profile_key": "general_business",
                "language": "en",
            },
        )
        assert status == 401
        assert payload["error"] == "browser_pairing_required"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)
