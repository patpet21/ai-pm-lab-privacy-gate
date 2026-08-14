from pathlib import Path
from copy import copy

from docx import Document
from openpyxl import Workbook, load_workbook
from reportlab.pdfgen import canvas

from ai_pm_lab_privacy_gate.domain.models import ReplacementMapping
from ai_pm_lab_privacy_gate.infrastructure.documents.restore_service import DocumentRestoreService


MAPPINGS = (
    ReplacementMapping("[[PG_PERSON_001]]", "PERSON", "Pietro Forestieri"),
    ReplacementMapping("[[PG_EMAIL_ADDRESS_001]]", "EMAIL_ADDRESS", "pietro@example.com"),
)


def test_restores_txt_and_reports_unknown_token(tmp_path: Path) -> None:
    source = tmp_path / "analysis.txt"
    source.write_text(
        "Owner: [[PG_PERSON_001]]\nUnknown: [[PG_PHONE_NUMBER_999]]",
        encoding="utf-8",
    )
    report = DocumentRestoreService().restore(source, MAPPINGS, tmp_path / "restored.txt")
    assert report.output_path.read_text(encoding="utf-8").startswith("Owner: Pietro Forestieri")
    assert report.restored_occurrences == 1
    assert report.unknown_tokens == ("[[PG_PHONE_NUMBER_999]]",)


def test_restores_docx_while_retaining_paragraph_style(tmp_path: Path) -> None:
    source = tmp_path / "analysis.docx"
    document = Document()
    paragraph = document.add_heading("Report for [[PG_PERSON_001]]", level=1)
    paragraph.runs[0].bold = True
    document.save(source)

    report = DocumentRestoreService().restore(source, MAPPINGS, tmp_path / "restored.docx")
    restored = Document(report.output_path)
    assert restored.paragraphs[0].text == "Report for Pietro Forestieri"
    assert restored.paragraphs[0].style.name == "Heading 1"
    assert restored.paragraphs[0].runs[0].bold


def test_restores_xlsx_while_retaining_sheet_style_and_formula(tmp_path: Path) -> None:
    source = tmp_path / "analysis.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Analysis"
    sheet["A1"] = "Owner"
    heading_font = copy(sheet["A1"].font)
    heading_font.bold = True
    sheet["A1"].font = heading_font
    sheet["B1"] = "[[PG_PERSON_001]]"
    sheet["C1"] = "=1+1"
    workbook.save(source)
    workbook.close()

    report = DocumentRestoreService().restore(source, MAPPINGS, tmp_path / "restored.xlsx")
    restored = load_workbook(report.output_path, data_only=False)
    try:
        assert restored["Analysis"]["B1"].value == "Pietro Forestieri"
        assert restored["Analysis"]["A1"].font.bold
        assert restored["Analysis"]["C1"].value == "=1+1"
    finally:
        restored.close()


def test_restores_searchable_pdf_to_layout_preserving_copy(tmp_path: Path) -> None:
    source = tmp_path / "analysis.pdf"
    pdf = canvas.Canvas(str(source))
    pdf.drawString(72, 720, "Report for [[PG_PERSON_001]]")
    pdf.save()

    report = DocumentRestoreService().restore(source, MAPPINGS, tmp_path / "restored.pdf")
    assert report.output_path.is_file()
    assert report.output_path.stat().st_size > 0
    assert report.restored_occurrences == 1


def test_repeated_restore_uses_a_new_output_when_preview_path_already_exists(tmp_path: Path) -> None:
    source = tmp_path / "analysis.pdf"
    pdf = canvas.Canvas(str(source))
    pdf.drawString(72, 720, "Report for [[PG_PERSON_001]]")
    pdf.save()

    service = DocumentRestoreService()
    destination = tmp_path / "restored.pdf"
    first = service.restore(source, MAPPINGS, destination)
    second = service.restore(source, MAPPINGS, destination)

    assert first.output_path == destination
    assert second.output_path != destination
    assert second.output_path.parent == destination.parent
    assert second.output_path.suffix == ".pdf"
    assert second.output_path.is_file()
    assert second.restored_occurrences == 1
