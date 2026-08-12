from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from ai_pm_lab_privacy_gate.domain.models import (
    Finding,
    PageContent,
    ProtectedSpan,
    ProtectionResult,
)
from ai_pm_lab_privacy_gate.infrastructure.documents.office_service import (
    OfficeDocumentService,
)


def _result_for(document, page_number: int, original: str, replacement: str) -> ProtectionResult:
    page = next(page for page in document.pages if page.page_number == page_number)
    start = page.text.index(original)
    finding = Finding(
        finding_id="test-finding",
        entity_type="EMAIL_ADDRESS",
        text=original,
        start=start,
        end=start + len(original),
        score=1.0,
        page_number=page_number,
        context=page.text,
    )
    protected_text = page.text[:start] + replacement + page.text[start + len(original):]
    protected_pages = tuple(
        PageContent(
            page_number=item.page_number,
            text=protected_text if item.page_number == page_number else item.text,
            location=item.location,
        )
        for item in document.pages
    )
    return ProtectionResult(
        protected_pages=protected_pages,
        applied_findings=(finding,),
        protected_spans=(
            ProtectedSpan(
                page_number=page_number,
                start=start,
                end=start + len(replacement),
                entity_type="EMAIL_ADDRESS",
                finding_id=finding.finding_id,
                replacement_text=replacement,
            ),
        ),
    )


def test_docx_round_trip_preserves_structure_and_replaces_selected_value(tmp_path: Path) -> None:
    source = tmp_path / "lease.docx"
    output = tmp_path / "lease_protected.docx"
    word = Document()
    paragraph = word.add_paragraph()
    paragraph.add_run("Tenant: ").bold = True
    paragraph.add_run("jane.smith@example.com")
    table = word.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Property 014"
    word.sections[0].header.paragraphs[0].text = "Private lease"
    word.save(source)

    service = OfficeDocumentService()
    document = service.extract(source)
    page = next(item for item in document.pages if "jane.smith@example.com" in item.text)
    result = _result_for(
        document, page.page_number, "jane.smith@example.com",
        "[[PG_EMAIL_ADDRESS_001]]",
    )
    service.write_protected(document, result, output)

    protected = Document(output)
    assert "jane.smith@example.com" not in "\n".join(
        item.text for item in protected.paragraphs
    )
    assert "[[PG_EMAIL_ADDRESS_001]]" in protected.paragraphs[0].text
    assert protected.paragraphs[0].runs[0].bold is True
    assert protected.tables[0].cell(0, 0).text == "Property 014"


def test_xlsx_round_trip_preserves_workbook_and_replaces_selected_cell(tmp_path: Path) -> None:
    source = tmp_path / "owners.xlsx"
    output = tmp_path / "owners_protected.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Owners"
    sheet["A1"] = "Email"
    sheet["B1"] = "jane.smith@example.com"
    sheet["A2"] = "Units"
    sheet["B2"] = 12
    sheet["B1"].number_format = "@"
    workbook.save(source)

    service = OfficeDocumentService()
    document = service.extract(source)
    page = next(item for item in document.pages if item.location == "Owners!B1")
    result = _result_for(
        document, page.page_number, "jane.smith@example.com",
        "[[PG_EMAIL_ADDRESS_001]]",
    )
    service.write_protected(document, result, output)

    protected = load_workbook(output, data_only=False)
    try:
        assert protected["Owners"]["B1"].value == "[[PG_EMAIL_ADDRESS_001]]"
        assert protected["Owners"]["B1"].number_format == "@"
        assert protected["Owners"]["B2"].value == 12
    finally:
        protected.close()
