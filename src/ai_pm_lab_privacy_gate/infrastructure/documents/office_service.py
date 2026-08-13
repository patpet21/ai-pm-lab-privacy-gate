from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, Literal

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell

from ai_pm_lab_privacy_gate.domain.models import AnalysisDocument, PageContent, ProtectionResult


OfficeKind = Literal["docx", "xlsx"]


@dataclass(frozen=True, slots=True)
class _DocxPart:
    paragraph: Paragraph


@dataclass(frozen=True, slots=True)
class _ExcelPart:
    cell: Cell
    component: Literal["value", "comment"]


class OfficeDocumentService:
    """Read and write protected Word and Excel copies without cloud services.

    Each editable Word paragraph or Excel value/comment becomes one analysis
    segment. The same deterministic traversal is used during export, so a
    Presidio result can be written back into a copy of the original document
    while retaining its layout, styles, tables, worksheets and formulas that do
    not contain selected sensitive values.
    """

    SUPPORTED_SUFFIXES = {".docx": "docx", ".xlsx": "xlsx"}

    def extract(self, path: str | Path) -> AnalysisDocument:
        source = self._validated_source(path)
        kind: OfficeKind = self.SUPPORTED_SUFFIXES[source.suffix.lower()]  # type: ignore[assignment]
        if kind == "docx":
            document = Document(str(source))
            entries = [
                (part.paragraph.text, "")
                for part in self._iter_docx_parts(document)
                if part.paragraph.text.strip()
            ]
        else:
            workbook = load_workbook(source, data_only=False, keep_links=False)
            try:
                entries = [
                    (self._excel_part_text(part), self._excel_part_location(part))
                    for part in self._iter_excel_parts(workbook)
                    if self._excel_part_text(part).strip()
                ]
            finally:
                workbook.close()
        pages = tuple(
            PageContent(
                page_number=index,
                text=text,
                location=location or f"Word block {index}",
            )
            for index, (text, location) in enumerate(entries, start=1)
        )
        return AnalysisDocument(source_kind=kind, source_path=source, pages=pages)

    def write_protected(
        self,
        source_document: AnalysisDocument,
        result: ProtectionResult,
        path: str | Path,
    ) -> Path:
        if (
            source_document.source_path is None
            or source_document.source_kind not in {"docx", "xlsx"}
        ):
            raise ValueError("A Word or Excel source document is required.")
        source = self._validated_source(source_document.source_path)
        destination = Path(path)
        expected_suffix = f".{source_document.source_kind}"
        if destination.suffix.lower() != expected_suffix:
            destination = destination.with_suffix(expected_suffix)
        destination.parent.mkdir(parents=True, exist_ok=True)
        if source_document.source_kind == "docx":
            self._write_docx(source, source_document, result, destination)
        else:
            self._write_xlsx(source, source_document, result, destination)
        return destination

    @classmethod
    def _validated_source(cls, path: str | Path) -> Path:
        source = Path(path)
        if not source.exists():
            raise FileNotFoundError(source)
        if source.suffix.lower() not in cls.SUPPORTED_SUFFIXES:
            raise ValueError("Supported Office formats are .docx and .xlsx.")
        return source

    def _write_docx(
        self,
        source: Path,
        source_document: AnalysisDocument,
        result: ProtectionResult,
        destination: Path,
    ) -> None:
        document = Document(str(source))
        parts = [
            part for part in self._iter_docx_parts(document)
            if part.paragraph.text.strip()
        ]
        self._verify_unchanged(
            source_document, [part.paragraph.text for part in parts]
        )
        replacements = self._replacements_by_page(result)
        for page_number, part in enumerate(parts, start=1):
            for start, end, replacement in reversed(
                replacements.get(page_number, ())
            ):
                self._replace_paragraph_range(
                    part.paragraph, start, end, replacement
                )
        self._clear_docx_identity(document)
        document.save(str(destination))

    def _write_xlsx(
        self,
        source: Path,
        source_document: AnalysisDocument,
        result: ProtectionResult,
        destination: Path,
    ) -> None:
        workbook = load_workbook(source, data_only=False, keep_links=False)
        try:
            parts = [
                part for part in self._iter_excel_parts(workbook)
                if self._excel_part_text(part).strip()
            ]
            self._verify_unchanged(
                source_document, [self._excel_part_text(part) for part in parts]
            )
            protected_by_page = {
                page.page_number: page.text for page in result.protected_pages
            }
            selected_pages = {
                finding.page_number for finding in result.applied_findings
            }
            for page_number, part in enumerate(parts, start=1):
                if page_number not in selected_pages:
                    continue
                protected = protected_by_page[page_number]
                if part.component == "comment":
                    if part.cell.comment is not None:
                        part.cell.comment.text = protected
                else:
                    # If PII appears inside a formula, privacy takes priority
                    # over keeping that formula executable in the safe copy.
                    part.cell.value = protected
            self._clear_xlsx_identity(workbook)
            workbook.save(destination)
        finally:
            workbook.close()

    @staticmethod
    def _replacements_by_page(
        result: ProtectionResult,
    ) -> dict[int, list[tuple[int, int, str]]]:
        replacement_by_finding = {
            span.finding_id: span.replacement_text
            for span in result.protected_spans
        }
        by_page: dict[int, list[tuple[int, int, str]]] = {}
        for finding in result.applied_findings:
            replacement = replacement_by_finding.get(finding.finding_id)
            if replacement is not None:
                by_page.setdefault(finding.page_number, []).append(
                    (finding.start, finding.end, replacement)
                )
        return by_page

    @staticmethod
    def _verify_unchanged(
        source_document: AnalysisDocument, current_texts: list[str]
    ) -> None:
        if current_texts != [page.text for page in source_document.pages]:
            raise ValueError(
                "The source document changed after it was scanned. "
                "Scan it again before exporting."
            )

    @classmethod
    def _iter_docx_parts(
        cls, document: DocxDocument
    ) -> Iterator[_DocxPart]:
        for child in document.iter_inner_content():
            if isinstance(child, Paragraph):
                yield _DocxPart(child)
            elif isinstance(child, Table):
                yield from cls._iter_table_parts(child)
        seen: set[int] = set()
        for section in document.sections:
            containers = (
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            )
            for container in containers:
                marker = id(container._element)
                if marker in seen:
                    continue
                seen.add(marker)
                for paragraph in container.paragraphs:
                    yield _DocxPart(paragraph)
                for table in container.tables:
                    yield from cls._iter_table_parts(table)

    @classmethod
    def _iter_table_parts(cls, table: Table) -> Iterator[_DocxPart]:
        seen_cells: set[object] = set()
        for row in table.rows:
            for cell in row.cells:
                marker = cell._tc
                if marker in seen_cells:
                    continue
                seen_cells.add(marker)
                yield from cls._iter_cell_parts(cell)

    @classmethod
    def _iter_cell_parts(cls, cell: _Cell) -> Iterator[_DocxPart]:
        for child in cell.iter_inner_content():
            if isinstance(child, Paragraph):
                yield _DocxPart(child)
            elif isinstance(child, Table):
                yield from cls._iter_table_parts(child)

    @staticmethod
    def _replace_paragraph_range(
        paragraph: Paragraph, start: int, end: int, replacement: str
    ) -> None:
        runs = list(paragraph.runs)
        if not runs or "".join(run.text for run in runs) != paragraph.text:
            text = paragraph.text
            paragraph.text = text[:start] + replacement + text[end:]
            return
        cursor = 0
        start_run = start_offset = end_run = end_offset = None
        for index, run in enumerate(runs):
            next_cursor = cursor + len(run.text)
            if start_run is None and (
                start < next_cursor
                or (start == next_cursor and index == len(runs) - 1)
            ):
                start_run, start_offset = index, start - cursor
            if end_run is None and end <= next_cursor:
                end_run, end_offset = index, end - cursor
                break
            cursor = next_cursor
        if (
            start_run is None
            or end_run is None
            or start_offset is None
            or end_offset is None
        ):
            raise ValueError(
                "Unable to map a protected value back to the Word document."
            )
        if start_run == end_run:
            run = runs[start_run]
            run.text = (
                run.text[:start_offset] + replacement + run.text[end_offset:]
            )
            return
        runs[start_run].text = runs[start_run].text[:start_offset] + replacement
        for index in range(start_run + 1, end_run):
            runs[index].text = ""
        runs[end_run].text = runs[end_run].text[end_offset:]

    @staticmethod
    def _iter_excel_parts(workbook) -> Iterator[_ExcelPart]:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.value is not None and str(cell.value).strip():
                        yield _ExcelPart(cell, "value")
                    if cell.comment is not None and cell.comment.text.strip():
                        yield _ExcelPart(cell, "comment")

    @staticmethod
    def _excel_part_text(part: _ExcelPart) -> str:
        if part.component == "comment":
            return part.cell.comment.text if part.cell.comment is not None else ""
        return str(part.cell.value)

    @staticmethod
    def _excel_part_location(part: _ExcelPart) -> str:
        suffix = " comment" if part.component == "comment" else ""
        return f"{part.cell.parent.title}!{part.cell.coordinate}{suffix}"

    @staticmethod
    def _clear_docx_identity(document: DocxDocument) -> None:
        properties = document.core_properties
        properties.author = "AI PM LAB Privacy Gate"
        properties.last_modified_by = "AI PM LAB Privacy Gate"
        properties.comments = "Protected locally by AI PM LAB Privacy Gate"

    @staticmethod
    def _clear_xlsx_identity(workbook) -> None:
        workbook.properties.creator = "AI PM LAB Privacy Gate"
        workbook.properties.lastModifiedBy = "AI PM LAB Privacy Gate"
        workbook.properties.description = (
            "Protected locally by AI PM LAB Privacy Gate"
        )
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.comment is not None:
                        cell.comment.author = "AI PM LAB Privacy Gate"
