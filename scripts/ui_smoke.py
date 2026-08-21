from __future__ import annotations

import os
import tempfile
import uuid
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
os.environ.setdefault(
    "PRIVACY_GATE_DATA_DIR",
    str((Path(tempfile.gettempdir()) / "privacy-gate-ui-smoke" / uuid.uuid4().hex).resolve()),
)

from PySide6.QtWidgets import QApplication, QBoxLayout
from PySide6.QtTest import QTest
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from ai_pm_lab_privacy_gate.application.privacy_service import PrivacyGateService
from ai_pm_lab_privacy_gate.domain.models import PageContent
from ai_pm_lab_privacy_gate.infrastructure.documents.pdf_service import PdfDocumentService
from ai_pm_lab_privacy_gate.ui.main_window import MainWindow
from ai_pm_lab_privacy_gate.ui.fonts import install_app_font
from ai_pm_lab_privacy_gate.ui.styles import APP_STYLE


def wait_until(predicate, timeout_ms: int = 10_000) -> bool:
    """Process Qt events until an asynchronous UI result is ready."""
    elapsed = 0
    while elapsed < timeout_ms:
        QApplication.processEvents()
        if predicate():
            return True
        QTest.qWait(100)
        elapsed += 100
    QApplication.processEvents()
    return bool(predicate())


def main() -> int:
    output_dir = Path(tempfile.gettempdir()) / "privacy-gate-ui-smoke-output"
    output = output_dir / "privacy_gate_main.png"
    setup_output = output_dir / "privacy_gate_compact_setup.png"
    collapsed_output = output_dir / "privacy_gate_collapsed.png"
    library_output = output_dir / "privacy_gate_library.png"
    contact_output = output_dir / "privacy_gate_contact.png"
    mask_output = output_dir / "privacy_gate_mask_colors.png"
    pdf_output = output_dir / "privacy_gate_pdf_comparison.png"
    focus_output = output_dir / "privacy_gate_focus_preview.png"
    word_output = output_dir / "privacy_gate_word_comparison.png"
    excel_output = output_dir / "privacy_gate_excel_comparison.png"
    restore_output = output_dir / "privacy_gate_restore_comparison.png"
    compact_restore_output = output_dir / "privacy_gate_restore_compact_window.png"
    output.parent.mkdir(parents=True, exist_ok=True)
    app = QApplication([])
    install_app_font(app)
    app.setStyleSheet(APP_STYLE)
    service = PrivacyGateService()
    document = service.document_from_text(
        "Tenant Jane Smith can be reached at jane.smith@example.com or 212-555-5555. "
        "Social Security Number: 219-09-9999."
    )
    window = MainWindow(service=service)
    page = window.protection_page
    findings = service.analyze(document, page._current_profile())
    page.text_input.setPlainText(document.pages[0].text)
    page._analysis_ready((document, findings))
    window.show()
    app.processEvents()
    page.setup_toggle.setChecked(True)
    app.processEvents()
    if not window.grab().save(str(setup_output)):
        raise RuntimeError("Unable to save compact setup screenshot")
    page.setup_toggle.setChecked(False)
    if not window.grab().save(str(output)):
        raise RuntimeError("Unable to save UI screenshot")
    page.mode_combo.setCurrentIndex(page.mode_combo.findData("mask"))
    app.processEvents()
    # The redesigned flow deliberately separates detection from protection.
    # Changing the mode invalidates any previous result until the user confirms
    # the selection with the explicit Protect action.
    page._redesign_protect_button.click()
    app.processEvents()
    if not page.current_result or len(page.current_result.combined_spans) != len(findings):
        raise RuntimeError("Mask-mode color metadata is incomplete")
    if not window.grab().save(str(mask_output)):
        raise RuntimeError("Unable to save mask color screenshot")

    source_pdf = output_dir / "privacy_gate_source.pdf"
    PdfDocumentService().write_protected(
        (
            PageContent(1, "Tenant Jane Smith\nEmail jane.smith@example.com\nPhone 212-555-5555"),
            PageContent(2, "Social Security Number 219-09-9999"),
        ),
        source_pdf,
    )
    pdf_document = service.document_from_pdf(source_pdf)
    pdf_findings = service.analyze(pdf_document, page._current_profile())
    page.pdf_path.setText(str(source_pdf.resolve()))
    page.input_tabs.setCurrentIndex(1)
    page._analysis_ready((pdf_document, pdf_findings))
    page._redesign_protect_button.click()
    page.preview_tabs.setCurrentIndex(1)
    if not wait_until(
        lambda: page.original_pdf_document.pageCount() == 2
        and page.protected_pdf_document.pageCount() == 2,
        timeout_ms=30_000,
    ):
        raise RuntimeError(
            "PDF comparison did not load both two-page documents: "
            f"original={page.original_pdf_document.pageCount()} "
            f"protected={page.protected_pdf_document.pageCount()} "
            f"timer_active={page._pdf_preview_timer.isActive()} "
            f"operations={page._redesign_active_operations} "
            f"worker={page._redesign_preview_worker!r} "
            f"note={page.comparison_note.text()}"
        )
    page._finding_selected(0, 0)
    page.keep_this_button.click()
    page._redesign_protect_button.click()
    app.processEvents()
    if not page.current_result or len(page.current_result.applied_findings) != len(pdf_findings) - 1:
        raise RuntimeError("Keep original did not remove the selected item from protection")
    page.protect_this_button.click()
    page._redesign_protect_button.click()
    app.processEvents()
    if not page.current_result or len(page.current_result.applied_findings) != len(pdf_findings):
        raise RuntimeError("Protect this did not restore the selected item to protection")
    if not window.grab().save(str(pdf_output)):
        raise RuntimeError("Unable to save PDF comparison screenshot")
    page.focus_preview_button.setChecked(True)
    app.processEvents()
    if page.findings_card.isVisible() or page.setup_card.isVisible():
        raise RuntimeError("Full document view did not hide the review panels")
    if not window.grab().save(str(focus_output)):
        raise RuntimeError("Unable to save focused preview screenshot")
    page.focus_preview_button.setChecked(False)
    page.mode_combo.setCurrentIndex(page.mode_combo.findData("reversible"))

    source_docx = output_dir / "privacy_gate_source.docx"
    word = Document()
    word.add_heading("Property inspection report", level=1)
    word.add_paragraph("Prepared for Jane Smith at 125 Main Street, New York, NY 10001.")
    word_table = word.add_table(rows=2, cols=2)
    word_table.style = "Table Grid"
    word_table.cell(0, 0).text = "Tenant email"
    word_table.cell(0, 1).text = "jane.smith@example.com"
    word_table.cell(1, 0).text = "Phone"
    word_table.cell(1, 1).text = "212-555-5555"
    word.save(source_docx)
    word_document = service.document_from_file(source_docx)
    word_findings = service.analyze(word_document, page._current_profile())
    page.pdf_path.setText(str(source_docx.resolve()))
    page.input_tabs.setCurrentIndex(1)
    page._analysis_ready((word_document, word_findings))
    page._redesign_protect_button.click()
    if not wait_until(
        lambda: page.original_view_stack.currentIndex() == 1
        and page.original_office_view.tabs.count() == 1
    ):
        raise RuntimeError(
            "Built-in Word comparison did not load: "
            f"stack={page.original_view_stack.currentIndex()} tabs={page.original_office_view.tabs.count()} "
            f"note={page.comparison_note.text()}"
        )
    if not window.grab().save(str(word_output)):
        raise RuntimeError("Unable to save Word comparison screenshot")

    source_xlsx = output_dir / "privacy_gate_source.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Tenants"
    worksheet.append(["Tenant", "Email", "Phone", "Rent"])
    worksheet.append(["Jane Smith", "jane.smith@example.com", "212-555-5555", 2450])
    worksheet.append(["Robert Brown", "robert.brown@example.com", "646-555-0104", 2875])
    for cell in worksheet[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="0B7189")
    notes = workbook.create_sheet("Notes")
    notes["A1"] = "Property"
    notes["B1"] = "125 Main Street, New York, NY 10001"
    workbook.save(source_xlsx)
    workbook.close()
    excel_document = service.document_from_file(source_xlsx)
    excel_findings = service.analyze(excel_document, page._current_profile())
    page.pdf_path.setText(str(source_xlsx.resolve()))
    page._analysis_ready((excel_document, excel_findings))
    page._redesign_protect_button.click()
    if not wait_until(
        lambda: page.original_view_stack.currentIndex() == 1
        and page.original_office_view.tabs.count() == 2
    ):
        raise RuntimeError("Built-in Excel comparison did not load worksheet tabs")
    if not window.grab().save(str(excel_output)):
        raise RuntimeError("Unable to save Excel comparison screenshot")

    if page.current_result is None:
        raise RuntimeError("Excel protection result is unavailable for Restore smoke test")
    protected_xlsx = output_dir / "privacy_gate_ai_analysis_protected.xlsx"
    service.save_protected_office(page.current_result, protected_xlsx, source_document=excel_document)
    excel_saved = window.library.save(
        title="Tenant analysis protected",
        source_kind="xlsx",
        source_name=source_xlsx.name,
        profile_key="property_management",
        result=page.current_result,
        labels=("Restore test",),
    )
    restore_page = window.restore_page
    restore_page.refresh(excel_saved.document_id)
    restore_page._begin_load_file(protected_xlsx)
    if not wait_until(
        lambda: restore_page._active_worker is None
        and restore_page._source_path == protected_xlsx
    ):
        raise RuntimeError("Restore page did not finish loading the protected Excel file")
    restore_page._restore()
    if not wait_until(
        lambda: bool(restore_page._restored_path and restore_page._restored_path.exists())
    ):
        raise RuntimeError("Structured Excel restore did not create a restored file")
    if restore_page.output_office_view.tabs.count() != 2:
        raise RuntimeError("Restore Excel comparison did not load both worksheet tabs")
    window._show_page(2)
    app.processEvents()
    if not window.grab().save(str(restore_output)):
        raise RuntimeError("Unable to save Restore comparison screenshot")

    window.resize(1120, 920)
    app.processEvents()
    if window.sidebar.width() != 76:
        raise RuntimeError("Narrow windows must automatically collapse the sidebar")
    if restore_page.source_row.direction() != QBoxLayout.Direction.TopToBottom:
        raise RuntimeError("Restore inputs did not stack in the compact window")
    restore_button_right = restore_page.restore_button.mapTo(window, restore_page.restore_button.rect().topLeft()).x() + restore_page.restore_button.width()
    if restore_button_right > window.width():
        raise RuntimeError("Restore action moved outside the compact window")
    if not window.grab().save(str(compact_restore_output)):
        raise RuntimeError("Unable to save compact Restore screenshot")
    window.resize(1458, 920)
    app.processEvents()

    window._toggle_sidebar()
    app.processEvents()
    if window.sidebar.width() != 76:
        raise RuntimeError("Sidebar did not collapse to the expected width")
    if any(button.text() or button.icon().isNull() for button in window.nav_buttons):
        raise RuntimeError("Collapsed navigation must use recognizable icons without initials")
    if not window.grab().save(str(collapsed_output)):
        raise RuntimeError("Unable to save collapsed UI screenshot")
    result = service.protect(document, findings)
    saved = window.library.save(
        title="Lease 014 - Jane Smith",
        source_kind="text",
        source_name="Pasted text",
        profile_key="property_management",
        result=result,
        labels=("Lease", "Property 014"),
    )
    window.library.set_favorite(saved.document_id, True)
    window._toggle_sidebar()
    window._show_page(1)
    app.processEvents()
    if not window.library_page.backup_button.isEnabled():
        raise RuntimeError("Library backup action is unavailable")
    if not window.grab().save(str(library_output)):
        raise RuntimeError("Unable to save library UI screenshot")
    window._show_page(5)
    app.processEvents()
    if window.contact_page.message_input.height() > 96:
        raise RuntimeError("Contact form is not using the compact layout")
    if not window.grab().save(str(contact_output)):
        raise RuntimeError("Unable to save Contact screenshot")
    print(
        f"UI_OK {setup_output.resolve()} {output.resolve()} {mask_output.resolve()} {pdf_output.resolve()} "
        f"{focus_output.resolve()} {word_output.resolve()} {excel_output.resolve()} {restore_output.resolve()} "
        f"{compact_restore_output.resolve()} "
        f"{collapsed_output.resolve()} {library_output.resolve()} {contact_output.resolve()} "
        f"{len(findings)} findings sidebar={window.sidebar.width()}"
    )
    window.close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
